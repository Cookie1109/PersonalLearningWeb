from __future__ import annotations

import base64
import io
import logging
import re
from pathlib import Path
from typing import Any, Literal
from urllib.parse import urlparse

import httpx
import html2text
import requests
from docx import Document as DocxDocument
from lxml import html as lxml_html
from PyPDF2 import PdfReader

try:
    from readability import Document as ReadabilityDocument
except Exception:  # pragma: no cover - runtime fallback when optional dependency is unavailable
    ReadabilityDocument = None

try:
    from newspaper import Article as NewspaperArticle
except Exception:  # pragma: no cover - runtime fallback when optional dependency is unavailable
    NewspaperArticle = None

try:
    import trafilatura
except Exception:  # pragma: no cover - runtime fallback when optional dependency is unavailable
    trafilatura = None

from app.core.config import get_settings
from app.core.exceptions import AppException

logger = logging.getLogger("app.parser")

ParserSourceType = Literal["url", "pdf", "docx", "image"]

MAX_UPLOAD_FILE_BYTES = 15 * 1024 * 1024
MAX_URL_DOWNLOAD_BYTES = 5 * 1024 * 1024
MAX_EXTRACTED_TITLE_LENGTH = 180

SUPPORTED_IMAGE_MIME_TYPES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/webp",
}
SUPPORTED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
SUPPORTED_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
}
NOISE_HTML_TAGS = ("nav", "header", "footer", "aside", "script", "style", "form", "noscript", "iframe", "svg")
MEDIA_NOISE_TAGS = ("script", "style", "noscript", "iframe", "svg", "canvas", "video", "audio", "source", "picture", "img")


def _collapse_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _sanitize_extracted_line(value: str) -> str:
    normalized = value.replace("\xa0", " ").strip()
    if not normalized:
        return ""

    normalized = re.sub(r"[<>]+", " ", normalized)
    normalized = normalized.strip()
    if re.fullmatch(r"[-=_*~`|•·]+", normalized):
        return ""

    return normalized


def _normalize_extracted_text(value: str) -> str:
    normalized_input = (value or "").replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    for raw_line in normalized_input.split("\n"):
        cleaned = _sanitize_extracted_line(raw_line)
        lines.append(cleaned)

    normalized = "\n".join(lines).strip()
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized


def _normalize_extracted_title(value: str | None) -> str | None:
    if not value:
        return None

    normalized = _collapse_whitespace(value)
    normalized = normalized.strip("-_|:\u2014\u2013 ")
    if not normalized:
        return None

    if len(normalized) > MAX_EXTRACTED_TITLE_LENGTH:
        normalized = normalized[:MAX_EXTRACTED_TITLE_LENGTH].rstrip()
    return normalized or None


def _build_title_from_text_excerpt(text: str, *, max_chars: int = 40) -> str | None:
    normalized = _collapse_whitespace(text or "")
    if not normalized:
        return None
    if len(normalized) <= max_chars:
        return _normalize_extracted_title(normalized)
    return _normalize_extracted_title(f"{normalized[:max_chars].rstrip()}...")


def _extract_title_from_html(html_text: str) -> str | None:
    if not html_text.strip():
        return None

    if ReadabilityDocument is not None:
        try:
            readability_title = ReadabilityDocument(html_text).title()
            normalized_readability_title = _normalize_extracted_title(readability_title)
            if normalized_readability_title:
                return normalized_readability_title
        except Exception as exc:
            logger.debug("parser.readability_title_failed error=%s", str(exc))

    try:
        root = lxml_html.fromstring(html_text)
        title_nodes = root.xpath("//title/text()")
        if title_nodes:
            return _normalize_extracted_title(title_nodes[0])
    except Exception as exc:
        logger.debug("parser.lxml_title_failed error=%s", str(exc))

    return None


def _extract_title_from_file_name(file_name: str | None) -> str | None:
    stem = Path(file_name or "").stem
    return _normalize_extracted_title(stem)


def _normalize_model_name(raw_model: str) -> str:
    model = (raw_model or "").strip()
    if model.startswith("models/"):
        model = model.split("/", 1)[1]

    legacy_map = {
        "gemini-1.5-flash": "gemini-2.5-flash",
        "gemini-1.5-pro": "gemini-2.5-pro",
    }
    return legacy_map.get(model, model)


def _build_model_candidates(settings) -> list[str]:
    configured_flash_model = (settings.gemini_model or "").strip() or "gemini-2.5-flash"
    configured_pro_model = (settings.gemini_pro_model or "").strip() or "gemini-2.5-flash"

    candidates: list[str] = []
    for candidate in (
        configured_flash_model,
        _normalize_model_name(configured_flash_model),
        configured_pro_model,
        _normalize_model_name(configured_pro_model),
    ):
        if candidate and candidate not in candidates:
            candidates.append(candidate)

    return candidates


def _extract_gemini_text(payload: dict[str, Any]) -> str:
    candidates = payload.get("candidates")
    if not isinstance(candidates, list) or not candidates:
        return ""

    candidate = candidates[0]
    if not isinstance(candidate, dict):
        return ""

    content = candidate.get("content")
    if not isinstance(content, dict):
        return ""

    parts = content.get("parts")
    if not isinstance(parts, list):
        return ""

    chunks: list[str] = []
    for part in parts:
        if not isinstance(part, dict):
            continue
        text = part.get("text")
        if isinstance(text, str) and text.strip():
            chunks.append(text.strip())

    return "\n\n".join(chunks).strip()


def _strip_unstable_media_blocks(html_text: str) -> str:
    cleaned = html_text or ""
    for tag in MEDIA_NOISE_TAGS:
        cleaned = re.sub(rf"<{tag}[^>]*?>[\s\S]*?</{tag}>", " ", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(rf"<{tag}[^>]*/>", " ", cleaned, flags=re.IGNORECASE)
    return cleaned


def _convert_html_to_markdown(html_text: str) -> str:
    if not (html_text or "").strip():
        return ""

    converter = html2text.HTML2Text()
    converter.body_width = 0
    converter.ignore_images = True
    converter.ignore_emphasis = False
    converter.ignore_links = False
    converter.ignore_tables = False
    converter.wrap_links = False
    converter.single_line_break = False
    converter.skip_internal_links = True

    safe_html = _strip_unstable_media_blocks(html_text)
    try:
        markdown = converter.handle(safe_html)
    except Exception as exc:
        logger.debug("parser.markdown_convert_failed error=%s", str(exc))
        return ""

    return _normalize_extracted_text(markdown)


def _extract_trafilatura_text(html_text: str) -> str:
    if trafilatura is None:
        return ""

    try:
        extracted = trafilatura.extract(
            html_text,
            output_format="txt",
            include_links=False,
            include_images=False,
            include_tables=False,
            deduplicate=True,
            favor_precision=True,
        )
    except Exception as exc:
        logger.debug("parser.trafilatura_extract_failed error=%s", str(exc))
        return ""

    return _normalize_extracted_text(extracted or "")


def _extract_readability_text(html_text: str) -> str:
    if ReadabilityDocument is None:
        return ""

    try:
        summary_html = ReadabilityDocument(html_text).summary(html_partial=True)
    except Exception as exc:
        logger.debug("parser.readability_extract_failed error=%s", str(exc))
        return ""

    if not summary_html:
        return ""

    markdown_text = _convert_html_to_markdown(summary_html)
    if markdown_text:
        return markdown_text

    try:
        root = lxml_html.fromstring(summary_html)
        return _normalize_extracted_text(root.text_content())
    except Exception:
        return ""


def _extract_newspaper_text(*, url: str, html_text: str) -> str:
    if NewspaperArticle is None:
        return ""

    try:
        article = NewspaperArticle(url=url)
        article.set_html(html_text)
        article.parse()
    except Exception as exc:
        logger.debug("parser.newspaper_extract_failed url=%s error=%s", url, str(exc))
        return ""

    markdown_candidates = [
        _convert_html_to_markdown(getattr(article, "article_html", "") or ""),
        _normalize_extracted_text(getattr(article, "text", "") or ""),
    ]
    non_empty = [item for item in markdown_candidates if item]
    if not non_empty:
        return ""
    return max(non_empty, key=len)


def _extract_lxml_main_text(html_text: str) -> str:
    try:
        root = lxml_html.fromstring(html_text)
    except Exception:
        return ""

    for xpath in (
        "//nav",
        "//header",
        "//footer",
        "//aside",
        "//script",
        "//style",
        "//form",
        "//noscript",
        "//iframe",
    ):
        for node in root.xpath(xpath):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)

    candidates = root.xpath("//article|//main")
    if candidates:
        best = max(candidates, key=lambda item: len((item.text_content() or "").strip()))
        return _normalize_extracted_text(best.text_content())

    body_nodes = root.xpath("//body")
    if body_nodes:
        return _normalize_extracted_text(body_nodes[0].text_content())

    return _normalize_extracted_text(root.text_content())


def _validate_url(url: str) -> str:
    normalized_url = url.strip()
    parsed = urlparse(normalized_url)
    if parsed.scheme not in {"http", "https"}:
        raise AppException(
            status_code=400,
            message="URL must start with http or https",
            detail={"code": "PARSER_URL_INVALID"},
        )

    if not parsed.netloc:
        raise AppException(
            status_code=400,
            message="URL is invalid",
            detail={"code": "PARSER_URL_INVALID"},
        )

    return normalized_url


def _extract_text_from_html(html_text: str, *, url: str | None = None) -> str:
    trafilatura_text = _extract_trafilatura_text(html_text)
    readability_text = _extract_readability_text(html_text)
    newspaper_text = _extract_newspaper_text(url=url or "https://example.com", html_text=html_text) if url else ""
    lxml_text = _extract_lxml_main_text(html_text)

    ordered_candidates = [
        trafilatura_text,
        readability_text,
        newspaper_text,
        lxml_text,
    ]
    normalized = next((value for value in ordered_candidates if value), "")

    if not normalized:
        raise AppException(
            status_code=409,
            message="No readable text found from URL",
            detail={"code": "PARSER_TEXT_EMPTY"},
        )

    return normalized


def extract_text_from_url(*, url: str) -> tuple[str, str | None]:
    normalized_url = _validate_url(url)

    try:
        response = requests.get(
            normalized_url,
            timeout=20,
            headers={
                "User-Agent": "PersonalLearningParser/1.0",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            },
            allow_redirects=True,
        )
        response.raise_for_status()
    except requests.Timeout as exc:
        raise AppException(
            status_code=503,
            message="URL fetch timed out",
            detail={"code": "PARSER_URL_TIMEOUT"},
        ) from exc
    except requests.RequestException as exc:
        raise AppException(
            status_code=503,
            message="Unable to fetch URL",
            detail={"code": "PARSER_URL_FETCH_FAILED"},
        ) from exc

    content = response.content or b""
    if len(content) > MAX_URL_DOWNLOAD_BYTES:
        raise AppException(
            status_code=413,
            message="URL content is too large",
            detail={"code": "PARSER_URL_TOO_LARGE"},
        )

    content_type = (response.headers.get("content-type") or "").lower()
    if "text/html" not in content_type and "application/xhtml+xml" not in content_type and "text/plain" not in content_type:
        raise AppException(
            status_code=400,
            message="URL content type is not supported",
            detail={"code": "PARSER_URL_UNSUPPORTED_CONTENT"},
        )

    try:
        if "text/plain" in content_type:
            text = (response.text or "").strip()
            normalized = _normalize_extracted_text(text)
            if not normalized:
                raise AppException(
                    status_code=409,
                    message="No readable text found from URL",
                    detail={"code": "PARSER_TEXT_EMPTY"},
                )
            return normalized, _build_title_from_text_excerpt(normalized)

        html_text = response.text or ""
        extracted_text = _extract_text_from_html(html_text, url=normalized_url)
        extracted_title = _extract_title_from_html(html_text) or _build_title_from_text_excerpt(extracted_text)
        return extracted_text, extracted_title
    except AppException:
        raise
    except Exception as exc:
        logger.warning("parser.url_extract_failed url=%s error=%s", normalized_url, str(exc))
        raise AppException(
            status_code=503,
            message="Failed to extract text from URL",
            detail={"code": "PARSER_URL_EXTRACT_FAILED"},
        ) from exc


def extract_text_from_pdf_bytes(*, file_bytes: bytes) -> str:
    if not file_bytes:
        raise AppException(
            status_code=400,
            message="Uploaded file is empty",
            detail={"code": "PARSER_FILE_EMPTY"},
        )

    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        chunks: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text)

        normalized = _normalize_extracted_text("\n".join(chunks))
        if not normalized:
            raise AppException(
                status_code=409,
                message="No readable text found in PDF",
                detail={"code": "PARSER_TEXT_EMPTY"},
            )
        return normalized
    except AppException:
        raise
    except Exception as exc:
        raise AppException(
            status_code=409,
            message="Failed to read PDF file",
            detail={"code": "PARSER_PDF_READ_FAILED"},
        ) from exc


def extract_text_from_docx_bytes(*, file_bytes: bytes) -> str:
    if not file_bytes:
        raise AppException(
            status_code=400,
            message="Uploaded file is empty",
            detail={"code": "PARSER_FILE_EMPTY"},
        )

    try:
        document = DocxDocument(io.BytesIO(file_bytes))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        normalized = _normalize_extracted_text("\n".join(chunks))
        if not normalized:
            raise AppException(
                status_code=409,
                message="No readable text found in DOCX",
                detail={"code": "PARSER_TEXT_EMPTY"},
            )
        return normalized
    except AppException:
        raise
    except Exception as exc:
        raise AppException(
            status_code=409,
            message="Failed to read DOCX file",
            detail={"code": "PARSER_DOCX_READ_FAILED"},
        ) from exc


def extract_text_from_image_bytes_via_gemini(*, file_bytes: bytes, mime_type: str) -> str:
    if not file_bytes:
        raise AppException(
            status_code=400,
            message="Uploaded image is empty",
            detail={"code": "PARSER_FILE_EMPTY"},
        )

    settings = get_settings()
    api_key = (settings.gemini_api_key or "").strip()
    if not api_key:
        raise AppException(
            status_code=503,
            message="AI service is not configured",
            detail={"code": "LLM_API_KEY_MISSING"},
        )

    image_base64 = base64.b64encode(file_bytes).decode("utf-8")
    model_candidates = _build_model_candidates(settings)
    timeout_seconds = max(60.0, float(settings.gemini_timeout_seconds))

    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text": "Extract and return all readable text from this image accurately. Return only extracted text.",
                    },
                    {
                        "inlineData": {
                            "mimeType": mime_type,
                            "data": image_base64,
                        }
                    },
                ],
            }
        ],
        "generationConfig": {
            "temperature": 0.0,
            "maxOutputTokens": 4096,
        },
    }

    with httpx.Client(timeout=timeout_seconds) as client:
        for index, model_name in enumerate(model_candidates):
            endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"
            has_fallback = index < len(model_candidates) - 1

            try:
                response = client.post(endpoint, params={"key": api_key}, json=payload)
            except httpx.TimeoutException as exc:
                if has_fallback:
                    continue
                raise AppException(status_code=503, message="AI service timeout", detail={"code": "LLM_TIMEOUT"}) from exc
            except httpx.RequestError as exc:
                if has_fallback:
                    continue
                raise AppException(
                    status_code=503,
                    message="AI service network error",
                    detail={"code": "LLM_NETWORK_ERROR"},
                ) from exc

            if response.status_code in (401, 403):
                raise AppException(
                    status_code=503,
                    message="AI service authentication failed",
                    detail={"code": "LLM_AUTH_FAILED"},
                )

            if response.status_code >= 400:
                if has_fallback and response.status_code in (404, 429, 500, 503):
                    continue
                raise AppException(
                    status_code=503,
                    message="AI service unavailable",
                    detail={"code": "LLM_SERVICE_ERROR"},
                )

            try:
                response_payload = response.json()
            except ValueError as exc:
                if has_fallback:
                    continue
                raise AppException(
                    status_code=503,
                    message="AI service returned invalid response",
                    detail={"code": "LLM_INVALID_RESPONSE"},
                ) from exc

            extracted_text = _normalize_extracted_text(_extract_gemini_text(response_payload))
            if extracted_text:
                return extracted_text

            if has_fallback:
                continue
            raise AppException(
                status_code=409,
                message="No readable text found in image",
                detail={"code": "PARSER_TEXT_EMPTY"},
            )

    raise AppException(
        status_code=503,
        message="AI service unavailable",
        detail={"code": "LLM_SERVICE_ERROR"},
    )


def _detect_uploaded_source_type(*, file_name: str | None, content_type: str | None) -> ParserSourceType:
    suffix = Path(file_name or "").suffix.lower()
    normalized_content_type = (content_type or "").split(";")[0].strip().lower()

    if normalized_content_type == "application/pdf" or suffix == ".pdf":
        return "pdf"

    if normalized_content_type in SUPPORTED_DOCX_MIME_TYPES or suffix == ".docx":
        return "docx"

    if normalized_content_type in SUPPORTED_IMAGE_MIME_TYPES or suffix in SUPPORTED_IMAGE_EXTENSIONS:
        return "image"

    raise AppException(
        status_code=400,
        message="Unsupported file format",
        detail={"code": "PARSER_UNSUPPORTED_FORMAT"},
    )


def extract_text_from_uploaded_file(
    *,
    file_name: str | None,
    content_type: str | None,
    file_bytes: bytes,
) -> tuple[str, ParserSourceType, str | None, str | None]:
    if not file_bytes:
        raise AppException(
            status_code=400,
            message="Uploaded file is empty",
            detail={"code": "PARSER_FILE_EMPTY"},
        )

    if len(file_bytes) > MAX_UPLOAD_FILE_BYTES:
        raise AppException(
            status_code=413,
            message="Uploaded file is too large",
            detail={"code": "PARSER_FILE_TOO_LARGE"},
        )

    source_type = _detect_uploaded_source_type(file_name=file_name, content_type=content_type)
    normalized_content_type = (content_type or "").split(";")[0].strip().lower() or None
    extracted_title = _extract_title_from_file_name(file_name)

    if source_type == "pdf":
        return extract_text_from_pdf_bytes(file_bytes=file_bytes), "pdf", normalized_content_type, extracted_title

    if source_type == "docx":
        return extract_text_from_docx_bytes(file_bytes=file_bytes), "docx", normalized_content_type, extracted_title

    mime_type = normalized_content_type if normalized_content_type in SUPPORTED_IMAGE_MIME_TYPES else "image/jpeg"
    extracted_text = extract_text_from_image_bytes_via_gemini(file_bytes=file_bytes, mime_type=mime_type)
    image_title = extracted_title or _build_title_from_text_excerpt(extracted_text)
    return extracted_text, "image", mime_type, image_title
